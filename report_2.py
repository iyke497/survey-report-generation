from docx import Document

# Create a new Document for the Word report
doc = Document()

# Add a title to the document
doc.add_heading('Comprehensive Project Report', 0)

# Iterate through each sector and MDA to add enhanced narrative to the Word document
for sector, sector_group in excel_data.groupby(sector_column):
    for mda, mda_group in sector_group.groupby(mda_column):
        for _, row in mda_group.iterrows():
            # Add sector and MDA headings
            doc.add_heading(f"Sector: {sector}", level=1)
            doc.add_heading(f"MDA: {mda}", level=2)

            # General Information Section
            general_info = f"Title of Project: {row[project_title_column] if pd.notna(row[project_title_column]) else 'Not provided'}\n"
            general_info += f"{row['Description of Total Achievement Recorded Since Inception']}.\n"
            general_info += f"Start Date: {row['What is the project start date?']} - End Date: {row['What is the project end date?']}.\n"
            general_info += f"Project represents a significant effort to {'upgrade' if 'upgrade' in str(row['Deliverables']).lower() else 'develop'} the facilities and services of the {mda}, aiming to address critical needs and enhance its functionality.\n"

            # Location Information
            location_info = f"State: {row['State Located ?']}, Location Coordinates: {row['Where is this project located?']}\n"
            location_info += "Strategically situated to maximize impact, ensuring the benefits of this development are felt across the designated areas.\n"

            # Situational Analysis - Financial Performance
            project_cost = safe_format_number(row['What is the total project cost?'])
            appropriated_amount = safe_format_number(row['What is the total amount appropriated this year?'])
            percentage_completion = row['Percentage of Total Achievement Recorded Since Inception']
            financial_performance = f"Project Cost: {project_cost}, Appropriated Amount This Year: {appropriated_amount}, Percentage Completion: {percentage_completion}%.\n"
            financial_performance += "While budgetary constraints have influenced the project's pace, the progress made underscores the dedication to achieving its full potential.\n"

            # Situational Analysis - Results Delivery
            results_delivery = f"Deliverables: {row['Deliverables']}, Key Performance Indicators: {row['Key Performance Indicators']}.\n"
            results_delivery += "These accomplishments align with the project's core objectives, establishing a foundation for its intended outcomes and long-term benefits.\n"

            # Lessons Learned
            lessons_learned = f"Challenges: {row['What are the project challenges?']}, Status: {row['What is the project status?']}.\n"
            lessons_learned += "Lessons from the challenges faced highlight the importance of comprehensive planning and the need for flexibility in project execution to accommodate unforeseen hurdles.\n"

            # Recommendations
            recommendations = f"Recommendations: {row[recommendations_column]}.\n"
            recommendations += "Implementing these recommendations is essential for ensuring the project's sustainability and maximizing its impact on the target community.\n"

            # Add sections to the document
            doc.add_heading('General Information', level=3)
            doc.add_paragraph(general_info)

            doc.add_heading('Location Information', level=3)
            doc.add_paragraph(location_info)

            doc.add_heading('Situation Analysis', level=3)
            doc.add_heading('Financial Performance', level=4)
            doc.add_paragraph(financial_performance)

            doc.add_heading('Results Delivery', level=4)
            doc.add_paragraph(results_delivery)

            doc.add_heading('Lessons Learned', level=4)
            doc.add_paragraph(lessons_learned)

            doc.add_heading('Recommendations', level=3)
            doc.add_paragraph(recommendations)

            # Add a page break after each project's report
            doc.add_page_break()

# Save the document
output_word_path = '/mnt/data/enhanced_project_report.docx'
doc.save(output_word_path)

output_word_path
